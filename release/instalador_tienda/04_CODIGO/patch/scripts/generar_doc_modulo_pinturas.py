"""Genera documento Word con resumen del módulo Colores en tienda."""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
except ImportError:
    print('Instalando python-docx...')
    import subprocess

    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt

OUT = ROOT / 'docs' / 'modulos' / 'COLOR_EN_TIENDA_RESUMEN.docx'

CODE_FILES = [
    'blueprints/modulo_pinturas.py',
    'services/fabrica_color_service.py',
    'services/pintura_stock_palette_service.py',
    'services/modulo_pinturas_session_service.py',
    'services/pintura_cartilla_service.py',
    'templates/modulos/fabrica_color.html',
    'templates/modulos/pinturas_lab.html',
    'templates/modulos/includes/header_modulo_pinturas.html',
    'static/js/fabrica-color.js',
    'static/css/fabrica-color.css',
    'scripts/generar_mascaras_ambientes.py',
    'tests/test_modulo_pinturas.py',
]


def _h(doc, text, level=1):
    doc.add_heading(text, level=level)


def _p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para


def _bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')


def _code(doc, path: Path, content: str):
    _h(doc, str(path.relative_to(ROOT)).replace('\\', '/'), level=3)
    para = doc.add_paragraph()
    run = para.add_run(content)
    run.font.name = 'Consolas'
    run.font.size = Pt(7)


def build():
    doc = Document()
    title = doc.add_heading('Módulo Colores en tienda — LhexIA Ferretería SD', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _p(doc, f'Fecha: {date.today().isoformat()} · Repo: sistema_ventas_limpio')
    doc.add_paragraph()

    _h(doc, '1. Qué queremos (objetivo de producto)', 1)
    _bullet(doc, 'Asistente en mostrador / tablet / TV para que el cliente elija ambiente, color y cantidad de pintura.')
    _bullet(doc, 'Mostrar solo productos con stock real en tienda (ERP), no tintometría ni mezcla — eso es fase 2.')
    _bullet(doc, 'Vista previa del color en un ambiente (comedor, cocina, dormitorio, baño, living, fachada).')
    _bullet(doc, 'Cotización referencial (galones, complementos, total) y acciones: carrito / pedir en caja / WhatsApp.')
    _bullet(doc, 'Módulo aislado del catálogo vitrina público: se habilita desde caja con token o lab con preview flag.')
    _bullet(doc, 'Integración con Liz (asistente flotante), no barra intrusiva en el wizard.')
    _bullet(doc, 'Calidad visual comparable a Sodimac Fábrica de Color (referencia, no copia de marca).')

    _h(doc, '2. Problemas encontrados y estado', 1)
    problems = [
        ('404 al abrir lab', 'Faltaba VITRINA_FABRICA_COLOR_PREVIEW=1 y reinicio Flask.', 'Resuelto'),
        ('UI poco premium vs Sodimac', 'Grid de fotos grandes, workspace split, panel lateral.', 'Mejorado; iteraciones pendientes'),
        ('Fotos no correspondían al texto', 'URLs Pexels/Unsplash incorrectas o cache.', 'Corregido con fotos curadas por ambiente'),
        ('Color teñía toda la foto', 'Overlay CSS mix-blend-mode sobre rectángulo completo.', 'Reemplazado por canvas + máscaras PNG'),
        ('Tinte verde global', 'Blur gaussiano en máscara dejaba alpha > 0 en muebles.', 'Corregido: máscaras binarias sin blur'),
        ('Ácido muriático en paleta amarilla', 'Filtro es_pintura_con_color incompleto.', 'Exclusiones ampliadas en pintura_stock_palette_service'),
        ('Barra rosa Liz confusa', 'Duplicaba al botón flotante Liz.', 'Eliminada del wizard'),
        ('Visualizador vs Sodimac', 'Sodimac usa máscaras trazadas a mano / SDK por foto.', 'Gap conocido — ver sección 5'),
        ('Etiqueta Dormitorio con foto Living', 'dormitorio.jpg era foto de living.', 'Corregido'),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Problema'
    hdr[1].text = 'Causa'
    hdr[2].text = 'Estado'
    for prob, cause, status in problems:
        row = table.add_row().cells
        row[0].text = prob
        row[1].text = cause
        row[2].text = status

    _h(doc, '3. Arquitectura y rutas', 1)
    _h(doc, '3.1 Activación', 2)
    _bullet(doc, '.env.local: VITRINA_FABRICA_COLOR_PREVIEW=1 (lab)')
    _bullet(doc, 'Caja: POST /api/caja/modulo-pinturas/habilitar → token cliente')
    _bullet(doc, 'Lab: GET /modulos/pinturas/lab → GET /modulos/pinturas/lab/iniciar')
    _bullet(doc, 'Cliente: GET /modulos/pinturas/<token>')

    _h(doc, '3.2 API', 2)
    _bullet(doc, 'POST /api/modulos/pinturas/<token>/cotizar')
    _bullet(doc, 'POST /api/modulos/pinturas/<token>/liz-tip')

    _h(doc, '3.3 Archivos principales', 2)
    for rel in CODE_FILES:
        _bullet(doc, rel)

    _h(doc, '3.4 Assets', 2)
    _bullet(doc, 'static/img/fabrica-color/ambientes/*.jpg — fotos por ambiente')
    _bullet(doc, 'static/img/fabrica-color/masks/*.png — máscaras de muro (generar con scripts/generar_mascaras_ambientes.py)')
    _bullet(doc, 'data/pintura_cartilla_sd.json — cartilla fase 2 (FABRICA_COLOR_FUENTE=cartilla)')

    _h(doc, '3.5 Flujo wizard (4 pasos)', 2)
    _bullet(doc, 'Paso 1: grid grande de ambientes → auto avanza al paso 2')
    _bullet(doc, 'Paso 2: paleta stock ERP + brillo + preview canvas')
    _bullet(doc, 'Paso 3: m² → galones estimados')
    _bullet(doc, 'Paso 4: resumen, complementos, pedir en caja / carrito')

    _h(doc, '4. Variables de entorno', 1)
    _bullet(doc, 'VITRINA_FABRICA_COLOR_PREVIEW=1 — habilita lab')
    _bullet(doc, 'FABRICA_COLOR_FUENTE=stock_erp (default) | cartilla (fase 2 tintometría)')

    _h(doc, '5. Gap vs Sodimac y próximos pasos', 1)
    _p(doc, 'Sodimac usa fotos propias + máscaras de muro por píxel (SDK proveedor pintura). '
           'Nosotros usamos stock genérico + polígonos/máscaras aproximadas. '
           'El color solo pinta bien la franja superior del muro; no segmenta detrás del sofá como Sodimac.')
    _h(doc, 'Para empatar calidad visual (recomendado post SD-1 piso):', 2)
    _bullet(doc, '6 fotos reales o renders alineados a Santo Domingo')
    _bullet(doc, '6 máscaras PNG trazadas a mano (Photoshop/Figma) por foto')
    _bullet(doc, 'Regenerar: python scripts/generar_mascaras_ambientes.py')
    _bullet(doc, 'Opcional fase 2: FABRICA_COLOR_FUENTE=cartilla + tintometría Kölor/Topex')

    _h(doc, '6. Comandos útiles', 1)
    _bullet(doc, 'pytest tests/test_modulo_pinturas.py -v')
    _bullet(doc, 'python scripts/generar_mascaras_ambientes.py')
    _bullet(doc, 'python scripts/generar_doc_modulo_pinturas.py')

    doc.add_page_break()
    _h(doc, 'Anexo A — Código fuente completo', 1)
    _p(doc, 'Snapshot al generar este documento. Archivos en orden de dependencia.')

    for rel in CODE_FILES:
        path = ROOT / rel
        if not path.exists():
            _h(doc, f'{rel} (NO ENCONTRADO)', level=3)
            continue
        try:
            content = path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            content = path.read_text(encoding='latin-1', errors='replace')
        if len(content) > 120000:
            content = content[:120000] + '\n\n... [TRUNCADO — archivo muy largo] ...'
        _code(doc, path, content)
        doc.add_paragraph()

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f'OK -> {OUT}')


if __name__ == '__main__':
    build()
